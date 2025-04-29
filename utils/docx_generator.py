import io
import os
import pandas as pd
import numpy as np
import datetime
import tempfile
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docxtpl import DocxTemplate, InlineImage
from docx.enum.dml import MSO_THEME_COLOR_INDEX

class DocxGenerator:
    """
    Class for generating DOCX reports with formula support
    """
    
    def __init__(self, form_data, data, insights, missing_fig=None, numeric_fig=None, categorical_fig=None):
        """
        Initialize the DocxGenerator
        
        Args:
            form_data (dict): Form data submitted by user
            data (pd.DataFrame): The dataset being analyzed
            insights (list): List of insights generated from the data
            missing_fig (plotly.graph_objects.Figure, optional): Missing values chart
            numeric_fig (plotly.graph_objects.Figure, optional): Numeric distribution chart
            categorical_fig (plotly.graph_objects.Figure, optional): Categorical distribution chart
        """
        self.form_data = form_data
        self.data = data
        self.insights = insights if insights else []
        self.missing_fig = missing_fig
        self.numeric_fig = numeric_fig
        self.categorical_fig = categorical_fig
        # Initialize temp_images as instance variable to avoid scoping issues
        self.temp_images = []
        
    def _save_figure_as_image(self, fig):
        """
        Save a Plotly figure as an image
        
        Args:
            fig (plotly.graph_objects.Figure): The figure to save
        
        Returns:
            str: Path to the saved image
        """
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as temp:
            fig.write_image(temp.name, format='png', width=800, height=500)
            return temp.name
            
    def _insert_formula(self, document, formula_text, paragraph=None, is_dax=False):
        """
        Insert a mathematical formula into the document
        
        Args:
            document (Document): The document object
            formula_text (str): The formula text in LaTeX format or DAX format
            paragraph (Paragraph, optional): Paragraph to add the formula to. If None, creates a new paragraph.
            is_dax (bool, optional): If True, treat formula_text as DAX formula instead of LaTeX
            
        Returns:
            Paragraph: The paragraph containing the formula
        """
        if paragraph is None:
            paragraph = document.add_paragraph()
        
        if is_dax:
            # For DAX formulas, we use a different style (code-like formatting)
            run = paragraph.add_run(formula_text)
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 139)  # Dark blue color for DAX code
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            # Add a light gray shading to the paragraph for DAX formulas
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'F2F2F2')  # Light gray background
            paragraph._element.get_or_add_pPr().append(shading_elm)
        else:
            # Simpler approach for math formulas - just display as formatted text
            # instead of using Office Math Markup Language which can cause XML issues
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(formula_text)
            run.font.italic = True
            run.font.name = 'Cambria Math'
            run.font.size = Pt(12)
            run.font.bold = False
            
            # Add a note that this is a mathematical formula
            note_run = paragraph.add_run()
            note_run.add_break()
            note_run.add_text("(Mathematical Formula)")
            note_run.font.size = Pt(8)
            note_run.font.italic = True
            note_run.font.color.rgb = RGBColor(100, 100, 100)
        
        return paragraph
        
    def _insert_lookup_function(self, document, function_type, description=None):
        """
        Insert Excel/DAX lookup function example with explanation
        
        Args:
            document (Document): The document object
            function_type (str): Type of lookup function ('vlookup', 'hlookup', 'xlookup', or 'dax')
            description (str, optional): Custom description for the function
            
        Returns:
            None
        """
        # Add a heading for the lookup function
        if function_type.lower() == 'vlookup':
            heading = document.add_heading('VLOOKUP Function', level=3)
            formula = 'VLOOKUP(lookup_value, table_array, col_index_num, [range_lookup])'
            if not description:
                description = ('The VLOOKUP function searches for a value in the leftmost column of a table, ' +
                              'and returns a value in the same row from a column you specify. ' +
                              'The "V" stands for vertical, meaning it searches down the first column.')
                
            example = 'VLOOKUP("Apple", A2:C10, 2, FALSE)'
            example_explanation = 'This example looks for "Apple" in the first column of range A2:C10 and returns the corresponding value from the 2nd column.'
            
        elif function_type.lower() == 'hlookup':
            heading = document.add_heading('HLOOKUP Function', level=3)
            formula = 'HLOOKUP(lookup_value, table_array, row_index_num, [range_lookup])'
            if not description:
                description = ('The HLOOKUP function searches for a value in the top row of a table, ' +
                              'and returns a value in the same column from a row you specify. ' +
                              'The "H" stands for horizontal, meaning it searches across the first row.')
                
            example = 'HLOOKUP("Q1", A1:D5, 3, FALSE)'
            example_explanation = 'This example looks for "Q1" in the first row of range A1:D5 and returns the corresponding value from the 3rd row.'
            
        elif function_type.lower() == 'xlookup':
            heading = document.add_heading('XLOOKUP Function', level=3)
            formula = 'XLOOKUP(lookup_value, lookup_array, return_array, [if_not_found], [match_mode], [search_mode])'
            if not description:
                description = ('The XLOOKUP function is the modern replacement for VLOOKUP. It searches for a match in a range or array, ' +
                              'and returns the corresponding value from another range or array. Unlike VLOOKUP, it can search in any direction ' +
                              'and return values from any direction.')
                
            example = 'XLOOKUP("Product123", A2:A100, C2:C100, "Not Found", 0, 1)'
            example_explanation = 'This example looks for "Product123" in range A2:A100 and returns the corresponding value from C2:C100. If not found, it returns "Not Found".'
            
        elif function_type.lower() == 'dax':
            heading = document.add_heading('DAX LOOKUPVALUE Function', level=3)
            formula = 'LOOKUPVALUE(result_columnName, search_columnName, search_value [, search_columnName, search_value] ...)'
            if not description:
                description = ('The DAX LOOKUPVALUE function returns a value from a table when a match is found. ' +
                              'Unlike VLOOKUP, it can use multiple search conditions and works directly with table relationships in Power BI and Analysis Services.')
                
            example = 'LOOKUPVALUE(Products[Price], Products[ProductID], "P-100")'
            example_explanation = 'This DAX example looks for the product with ID "P-100" and returns its price from the Products table.'
        
        # Add the function syntax with formatting
        p = document.add_paragraph()
        p.add_run("Syntax: ").bold = True
        code_run = p.add_run(formula)
        code_run.font.name = "Consolas"
        code_run.font.size = Pt(10)
        
        # Add the description
        document.add_paragraph(description)
        
        # Add example with formatting
        p = document.add_paragraph()
        p.add_run("Example: ").bold = True
        code_run = p.add_run(example)
        code_run.font.name = "Consolas"
        code_run.font.size = Pt(10)
        
        # Add example explanation
        document.add_paragraph(example_explanation)
        
        # If it's DAX, add a more complex example
        if function_type.lower() == 'dax':
            document.add_heading("Complex DAX Example", level=4)
            complex_example = """// Calculate total sales for current product
TotalSales = 
CALCULATE(
    SUM(Sales[Amount]),
    FILTER(
        ALL(Products),
        Products[Category] = SELECTEDVALUE(Products[Category])
    )
)"""
            self._insert_formula(document, complex_example, is_dax=True)
            document.add_paragraph("This DAX measure calculates the total sales amount filtered by the currently selected product category.")
        
        # Add a horizontal line for separation
        p = document.add_paragraph()
        p.paragraph_format.bottom_border.space = Pt(2)
        p.paragraph_format.bottom_border.width = 0.5

    def generate_docx(self):
        """
        Generate a DOCX report with formula support
        
        Returns:
            bytes: DOCX file as bytes
        """
        # Create a new document
        doc = Document()
        
        # Set up document properties
        doc.core_properties.title = "Easy AI Analytics Report"
        doc.core_properties.author = "Easy AI Analytics"
        doc.core_properties.comments = f"Report generated on {datetime.datetime.now().strftime('%Y-%m-%d')}"
        
        # Add title
        title = doc.add_heading('Easy AI Analytics Report', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add company and project info
        doc.add_heading(f"Company: {self.form_data.get('company_name', 'N/A')}", level=1)
        doc.add_heading(f"Project: {self.form_data.get('project_title', 'N/A')}", level=1)
        
        # Add date
        date_paragraph = doc.add_paragraph()
        date_paragraph.add_run(f"Report Date: {datetime.datetime.now().strftime('%Y-%m-%d')}").bold = True
        
        # Add business objectives
        doc.add_heading("Business Objectives", level=2)
        doc.add_paragraph(self.form_data.get('objectives', 'N/A'))
        
        # Add target audience
        doc.add_heading("Target Audience", level=2)
        doc.add_paragraph(self.form_data.get('target_audience', 'N/A'))
        
        # Add timeframe
        doc.add_heading("Analysis Timeframe", level=2)
        doc.add_paragraph(self.form_data.get('timeframe', 'N/A'))
        
        # Add dataset overview with error handling
        doc.add_heading("Dataset Overview", level=2)
        if self.data is not None and not self.data.empty:
            p = doc.add_paragraph()
            p.add_run(f"Rows: ").bold = True
            p.add_run(f"{self.data.shape[0]}")
            
            p = doc.add_paragraph()
            p.add_run(f"Columns: ").bold = True
            p.add_run(f"{self.data.shape[1]}")
            
            p = doc.add_paragraph()
            p.add_run(f"Missing Values: ").bold = True
            p.add_run(f"{self.data.isna().sum().sum()}")
            
            # Add a sample formula for dataset statistics
            doc.add_heading("Statistical Formulas Used in Analysis", level=3)
            self._insert_formula(doc, "\\bar{x} = \\frac{1}{n}\\sum_{i=1}^{n}x_i")
            doc.add_paragraph("Formula for calculating the mean value of the dataset")
            
            self._insert_formula(doc, "s = \\sqrt{\\frac{1}{n-1}\\sum_{i=1}^{n}(x_i - \\bar{x})^2}")
            doc.add_paragraph("Formula for calculating the standard deviation")
            
            # Add correlation formula if data includes numeric columns
            if self.data.select_dtypes(include=['number']).shape[1] > 1:
                self._insert_formula(doc, "r_{xy} = \\frac{\\sum_{i=1}^{n}(x_i-\\bar{x})(y_i-\\bar{y})}{\\sqrt{\\sum_{i=1}^{n}(x_i-\\bar{x})^2\\sum_{i=1}^{n}(y_i-\\bar{y})^2}}")
                doc.add_paragraph("Formula for calculating the Pearson correlation coefficient between variables")
                
            # Add Excel and DAX lookup functions reference section
            doc.add_heading("Excel and DAX Lookup Functions Reference", level=2)
            doc.add_paragraph("This section provides reference for Excel lookup functions and DAX query expressions commonly used in data analysis and reporting.")
            
            # Add VLOOKUP, HLOOKUP, and XLOOKUP examples
            self._insert_lookup_function(doc, 'vlookup')
            self._insert_lookup_function(doc, 'hlookup')
            self._insert_lookup_function(doc, 'xlookup')
            
            # Add DAX lookup example
            self._insert_lookup_function(doc, 'dax')
            
            # Add additional DAX formulas and examples section
            doc.add_heading("Additional DAX Formulas", level=2)
            
            # Time intelligence DAX formula
            doc.add_heading("DAX Time Intelligence", level=3)
            time_intelligence_dax = """// Sales Year-to-Date
Sales YTD = 
CALCULATE(
    SUM(Sales[Amount]),
    DATESYTD(Calendar[Date])
)

// Year-over-Year Growth
YOY Growth = 
VAR CurrentYearSales = [Total Sales]
VAR PreviousYearSales = 
    CALCULATE(
        [Total Sales],
        SAMEPERIODLASTYEAR(Calendar[Date])
    )
RETURN
    DIVIDE(CurrentYearSales - PreviousYearSales, PreviousYearSales)"""
            
            self._insert_formula(doc, time_intelligence_dax, is_dax=True)
            doc.add_paragraph("These DAX formulas calculate Year-to-Date sales and Year-over-Year growth using DAX time intelligence functions.")
            
            # Filter context DAX formula
            doc.add_heading("DAX Filter Context", level=3)
            filter_context_dax = """// Market Share Calculation
Market Share = 
DIVIDE(
    [Total Sales],
    CALCULATE(
        [Total Sales],
        ALL(Products)
    )
)

// Top 5 Products
Top 5 Products = 
IF(
    RANKX(
        ALL(Products),
        [Total Sales],
        ,
        DESC
    ) <= 5,
    [Total Sales],
    BLANK()
)"""
            
            self._insert_formula(doc, filter_context_dax, is_dax=True)
            doc.add_paragraph("These DAX measures demonstrate filter context manipulation for calculating market share and identifying top products by sales.")
        else:
            doc.add_paragraph("No data available for analysis")
        
        # Add visualizations if available
        if any([self.missing_fig, self.numeric_fig, self.categorical_fig]):
            doc.add_heading("Data Visualizations", level=2)
            
            if self.missing_fig:
                try:
                    img_path = self._save_figure_as_image(self.missing_fig)
                    doc.add_heading("Missing Values Chart", level=3)
                    doc.add_picture(img_path, width=Inches(6))
                    self.temp_images.append(img_path)
                except Exception as e:
                    doc.add_paragraph(f"Error generating missing values chart: {e}")
            
            if self.numeric_fig:
                try:
                    img_path = self._save_figure_as_image(self.numeric_fig)
                    doc.add_heading("Numeric Distribution Chart", level=3)
                    doc.add_picture(img_path, width=Inches(6))
                    self.temp_images.append(img_path)
                except Exception as e:
                    doc.add_paragraph(f"Error generating numeric distribution chart: {e}")
            
            if self.categorical_fig:
                try:
                    img_path = self._save_figure_as_image(self.categorical_fig)
                    doc.add_heading("Categorical Distribution Chart", level=3)
                    doc.add_picture(img_path, width=Inches(6))
                    self.temp_images.append(img_path)
                except Exception as e:
                    doc.add_paragraph(f"Error generating categorical distribution chart: {e}")
        
        # Add key metrics
        doc.add_heading("Key Metrics to Track", level=2)
        doc.add_paragraph(self.form_data.get('key_metrics', 'N/A'))
        
        # Add insights
        doc.add_heading("Data Insights", level=2)
        for insight in self.insights:
            doc.add_paragraph(insight, style='List Bullet')
        
        # Add sample data table
        doc.add_heading("Data Sample", level=2)
        
        if self.data is not None and not self.data.empty:
            # Convert any problematic columns to string type
            safe_data = self.data.copy()
            for col in safe_data.columns:
                if pd.api.types.is_datetime64_any_dtype(safe_data[col]):
                    safe_data[col] = safe_data[col].dt.strftime('%Y-%m-%d %H:%M:%S')
            
            # Create a table for the data
            # Add headers as first row
            table = doc.add_table(rows=1, cols=len(safe_data.columns))
            table.style = 'Table Grid'
            
            # Add header row
            header_cells = table.rows[0].cells
            for i, column_name in enumerate(safe_data.columns):
                header_cells[i].text = str(column_name)
                # Format header cells
                for paragraph in header_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
            
            # Add data rows (up to 5 rows)
            for i in range(min(5, len(safe_data))):
                row_cells = table.add_row().cells
                for j, value in enumerate(safe_data.iloc[i]):
                    if isinstance(value, (pd.Timestamp, np.datetime64)):
                        cell_text = str(pd.Timestamp(value).strftime('%Y-%m-%d %H:%M:%S'))
                    else:
                        cell_text = str(value)
                    row_cells[j].text = cell_text
        else:
            doc.add_paragraph("No data available to display")
        
        # Save to bytes
        f = io.BytesIO()
        doc.save(f)
        
        # Clean up temporary image files
        for img_path in self.temp_images:
            try:
                os.unlink(img_path)
            except:
                pass
        
        # Get the value of the buffer and return it
        f.seek(0)
        return f.getvalue()